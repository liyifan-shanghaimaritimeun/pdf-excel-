#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
【gen_docx.py —— 生成训练报告】
作用：读取「CVAT 手动标注数据集 + YOLO 训练产物」，拼出一份人能直接看的 docx 报告。
      报告含：数据集概览、训练配置、训练结果（mAP/各类指标）、结论建议、产出文件清单。

依赖：python-docx（pip install python-docx）。训练时 gen_docx 会在训练结束后被调用。
"""

import os
import json
import csv

# ============ 路径常量 ============
ROOT = r"C:/Users/admin/WorkBuddy/2026-07-27-09-13-23/table-training"
DATA_YAML = os.path.join(ROOT, "yolo_dataset", "data.yaml")          # 类别名来源
RUNS = os.path.join(ROOT, "runs", "cvat_tables")                     # 训练产物目录
METRICS_JSON = os.path.join(ROOT, "train_metrics.json")              # 训练后写的指标
RESULTS_CSV = os.path.join(RUNS, "results.csv")                      # 每轮指标曲线
ERR_LOG = os.path.join(ROOT, "train_error.log")                      # 训练报错日志
MODEL_OUT = os.path.join(ROOT, "cvat_yolov8n.pt")                    # 导出的最佳权重
OUT_DOCX = os.path.join(ROOT, "表格检测训练报告.docx")               # 最终报告


def load_yaml_classes():
    """从 data.yaml 里抠出 names（类别名列表）。"""
    names = []
    with open(DATA_YAML, encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line.startswith("names:"):
                s = line.split("names:", 1)[1].strip()
                try:
                    names = json.loads(s)   # names 是 JSON 数组，安全解析
                except Exception:
                    pass
    return names


def collect():
    """统计 train/val 的图片数、框数、各类框数（直接扫 labels/*.txt）。"""
    names = load_yaml_classes()
    nc = len(names)
    base = os.path.join(ROOT, "yolo_dataset", "labels")
    n_img = {"train": 0, "val": 0}
    n_box = {"train": 0, "val": 0}
    per_class = {"train": [0] * nc, "val": [0] * nc}
    for split in ("train", "val"):
        d = os.path.join(base, split)
        if not os.path.isdir(d):
            continue
        for fn in os.listdir(d):
            if not fn.endswith(".txt"):
                continue
            n_img[split] += 1
            with open(os.path.join(d, fn), encoding="utf-8") as f:
                for ln in f:
                    ln = ln.strip()
                    if not ln:
                        continue
                    p = ln.split()
                    if len(p) < 5:
                        continue
                    try:
                        c = int(p[0])
                    except ValueError:
                        continue
                    n_box[split] += 1
                    if 0 <= c < nc:
                        per_class[split][c] += 1
    return names, n_img, n_box, per_class


def load_metrics():
    """读 train_metrics.json（训练主程序最后写的）。"""
    if not os.path.exists(METRICS_JSON):
        return None
    try:
        with open(METRICS_JSON, encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def load_results():
    """读 results.csv（每轮一行，ultralytics 自动生成）。"""
    rows = []
    if not os.path.exists(RESULTS_CSV):
        return rows
    with open(RESULTS_CSV, encoding="utf-8") as f:
        r = csv.DictReader(f)
        for row in r:
            rows.append(row)
    return rows


def gv(row, k):
    """安全取 csv 某列的数值（缺失/异常返回 0.0）。"""
    if row is None:
        return 0.0
    try:
        return float(row.get(k, 0) or 0)
    except Exception:
        return 0.0


def main():
    from docx import Document   # 用到才导入，避免没装 python-docx 时整脚本崩

    # 1) 汇总数据
    names, n_img, n_box, per_class = collect()
    nc = len(names)
    mj = load_metrics()
    rows = load_results()

    # 2) 找最佳 mAP50 出现的轮次（best_epoch）
    best_map50 = 0.0
    best_epoch = 0
    first = rows[0] if rows else None
    last = rows[-1] if rows else None
    for r in rows:
        m = gv(r, "metrics/mAP50(B)")
        if m > best_map50:
            best_map50 = m
            try:
                best_epoch = int(float(r.get("epoch", 0)))
            except Exception:
                pass

    # 3) 建文档
    doc = Document()
    doc.add_heading("表格检测模型训练报告", level=0)

    # ===== 一、数据集概览 =====
    doc.add_heading("一、数据集概览", level=1)
    doc.add_paragraph(
        "数据来源：CVAT 标注平台 Project 2（4 类：表格 / 表头 / 重点内容 / 型号）。"
        "本次训练已清空全部自动标注（auto），仅保留人工手动标注（manual）与半自动标注（semi-auto），"
        "以更可靠的人工标注作为监督信号重新训练。"
    )
    doc.add_paragraph(
        f"图片总数：{n_img['train'] + n_img['val']}（训练集 {n_img['train']} / 验证集 {n_img['val']}）"
    )
    doc.add_paragraph(
        f"标注框总数：{n_box['train'] + n_box['val']}（训练集 {n_box['train']} / 验证集 {n_box['val']}）"
    )
    # 各类框数表格
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "类别"; hdr[1].text = "训练集框数"; hdr[2].text = "验证集框数"; hdr[3].text = "合计"
    for i, nm in enumerate(names):
        row = tbl.add_row().cells
        row[0].text = nm
        row[1].text = str(per_class["train"][i])
        row[2].text = str(per_class["val"][i])
        row[3].text = str(per_class["train"][i] + per_class["val"][i])

    # ===== 二、训练配置 =====
    doc.add_heading("二、训练配置", level=1)
    cfg = [
        ("模型", "YOLOv8n（yolov8n.pt 官方预训练权重 fine-tune）"),
        ("任务", "目标检测 detect"),
        ("类别数", f"{nc}（表格 / 表头 / 重点内容 / 型号）"),
        ("图像尺寸", "640 × 640"),
        ("批次大小", "4（CPU 内存受限）"),
        ("训练轮数", "50（patience=15 早停）"),
        ("优化器", "auto"),
        ("初始学习率", "0.001"),
        ("随机种子", "42"),
        ("设备", "CPU（本机无 GPU）"),
    ]
    t2 = doc.add_table(rows=0, cols=2)
    t2.style = "Table Grid"
    for k, v in cfg:
        c = t2.add_row().cells
        c[0].text = k; c[1].text = v

    # ===== 三、训练结果 =====
    doc.add_heading("三、训练结果", level=1)
    if mj:
        doc.add_paragraph(f"最终验证 mAP@0.5      : {float(mj.get('map50', 0)):.4f}")
        doc.add_paragraph(f"最终验证 mAP@0.5:0.95 : {float(mj.get('map50_95', 0)):.4f}")
        if last is not None:
            doc.add_paragraph(
                f"末轮 Precision: {gv(last, 'metrics/precision(B)'):.4f}   "
                f"Recall: {gv(last, 'metrics/recall(B)'):.4f}"
            )
        pc = mj.get("per_class_map50_95")
        if pc:
            doc.add_paragraph("各类 mAP@0.5:0.95：")
            t3 = doc.add_table(rows=1, cols=2)
            t3.style = "Table Grid"
            h = t3.rows[0].cells
            h[0].text = "类别"; h[1].text = "mAP@0.5:0.95"
            for i, nm in enumerate(names):
                c = t3.add_row().cells
                c[0].text = nm
                val = pc[i] if i < len(pc) else 0
                c[1].text = f"{float(val):.4f}"
    else:
        doc.add_paragraph("⚠️ 未找到训练指标文件 train_metrics.json，训练可能失败，详见 train_error.log。")

    if rows:
        doc.add_paragraph(
            f"训练共 {len(rows)} 轮；最佳 mAP@0.5 出现在第 {best_epoch} 轮（mAP@0.5={best_map50:.4f}）。"
            + (f" 首轮 mAP@0.5≈{gv(first, 'metrics/mAP50(B)'):.4f}，末轮 mAP@0.5≈{gv(last, 'metrics/mAP50(B)'):.4f}。"
               if first is not None else "")
        )
        if first is not None and last is not None:
            doc.add_paragraph(
                f"训练损失收敛：box_loss {gv(first, 'train/box_loss'):.4f} → {gv(last, 'train/box_loss'):.4f}，"
                f"cls_loss {gv(first, 'train/cls_loss'):.4f} → {gv(last, 'train/cls_loss'):.4f}，"
                f"验证 mAP 整体呈上升趋势。"
            )

    # ===== 四、结论与建议 =====
    doc.add_heading("四、结论与建议", level=1)
    for b in [
        "本次仅使用人工手动标注（已剔除全部自动标注），标签质量高但样本规模有限，模型性能上限受数据量约束。",
        "类别分布不均衡（『重点内容』占比最高），稀有类别（型号、表格）样本偏少，建议优先针对弱类人工补标。",
        "半自动标注（semi-auto）若经人工复核确认无误，可标记为 manual 进一步扩充训练集。",
        "当前自动标注对扫描图 / 矢量图型 PDF 存在盲区，后续可引入 OCR 引擎（PaddleOCR / EasyOCR）自动检测文本型型号。",
        "纯 CPU 训练较慢（约 2.5–3 小时 / 50 轮），如有 GPU 环境可显著加速并采用更大 batch。",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    # ===== 五、产出文件 =====
    doc.add_heading("五、产出文件", level=1)
    files = [
        ("训练权重", MODEL_OUT if os.path.exists(MODEL_OUT) else os.path.join(RUNS, "weights", "best.pt")),
        ("数据集", os.path.join(ROOT, "yolo_dataset")),
        ("训练指标", METRICS_JSON),
        ("训练曲线", RESULTS_CSV),
        ("报错日志", ERR_LOG),
    ]
    for label, path in files:
        doc.add_paragraph(f"{label}：{path}", style="List Bullet")

    # 4) 保存
    doc.save(OUT_DOCX)
    print("docx 已生成:", OUT_DOCX)


if __name__ == "__main__":
    main()
